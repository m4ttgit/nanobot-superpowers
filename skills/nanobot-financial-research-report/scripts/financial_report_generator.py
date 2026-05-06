#!/usr/bin/env python3
"""financial_report_generator.py â€” Generate a professional financial research report."""
import argparse
import json
import sys
from datetime import datetime


def calculate_all_ratios(data):
    """Calculate all 5 categories of financial ratios."""
    income = data.get('income_statement', {})
    balance = data.get('balance_sheet', {})

    revenue = abs(income.get('revenue', 0))
    cogs = abs(income.get('cogs', 0))
    gross_profit = abs(income.get('gross_profit', 0))
    ebit = abs(income.get('ebit', 0))
    ebitda = abs(income.get('ebitda', 0))
    net_income = abs(income.get('net_income', 0))

    current_assets = abs(balance.get('current_assets', 0))
    total_assets = abs(balance.get('total_assets', 0))
    inventory = abs(balance.get('inventory', 0))
    ar = abs(balance.get('accounts_receivable', 0))
    cash = abs(balance.get('cash', 0))
    current_liabilities = abs(balance.get('current_liabilities', 0))
    total_liabilities = abs(balance.get('total_liabilities', 0))
    total_equity = abs(balance.get('total_equity', 0))
    total_debt = abs(balance.get('total_debt', total_liabilities))

    result = {'profitability': {}, 'liquidity': {}, 'leverage': {}, 'efficiency': {}, 'valuation': {}}

    # Profitability
    if revenue > 0:
        result['profitability']['gross_margin'] = gross_profit / revenue
        result['profitability']['operating_margin'] = ebit / revenue
        result['profitability']['net_margin'] = net_income / revenue
    if total_assets > 0:
        result['profitability']['roa'] = net_income / total_assets
    if total_equity > 0:
        result['profitability']['roe'] = net_income / total_equity
    if revenue > 0:
        result['profitability']['ebitda_margin'] = ebitda / revenue

    # Liquidity
    if current_liabilities > 0:
        result['liquidity']['current_ratio'] = current_assets / current_liabilities
        result['liquidity']['quick_ratio'] = (current_assets - inventory) / current_liabilities
    result['liquidity']['working_capital'] = current_assets - current_liabilities
    if cash > 0 and current_liabilities > 0:
        result['liquidity']['cash_ratio'] = cash / current_liabilities

    # Leverage
    if total_equity > 0:
        result['leverage']['debt_to_equity'] = total_debt / total_equity
    if total_assets > 0:
        result['leverage']['debt_to_assets'] = total_debt / total_assets
    if ebit > 0:
        interest = abs(income.get('interest', 0))
        result['leverage']['interest_coverage'] = ebit / interest if interest > 0 else None
    if total_assets > 0 and total_equity > 0:
        result['leverage']['financial_leverage'] = total_assets / total_equity

    # Efficiency
    if total_assets > 0 and revenue > 0:
        result['efficiency']['asset_turnover'] = revenue / total_assets
    if cogs > 0 and inventory > 0:
        result['efficiency']['inventory_turnover'] = cogs / inventory
        result['efficiency']['inventory_days'] = (inventory / cogs) * 365
    if revenue > 0 and ar > 0:
        result['efficiency']['receivables_turnover'] = revenue / ar
        result['efficiency']['debtor_days'] = (ar / revenue) * 365
    if cogs > 0 and total_liabilities > 0:
        ap = abs(balance.get('accounts_payable', 0))
        result['efficiency']['creditor_days'] = (ap / cogs) * 365
    if all(v is not None for v in [result['efficiency'].get('debtor_days'), result['efficiency'].get('inventory_days'), result['efficiency'].get('creditor_days')]):
        result['efficiency']['cash_conversion_cycle'] = (
            result['efficiency']['debtor_days'] +
            result['efficiency']['inventory_days'] -
            result['efficiency']['creditor_days']
        )

    # Valuation (if market data provided)
    market_cap = data.get('market_cap')
    enterprise_value = data.get('enterprise_value')
    if market_cap and net_income > 0:
        result['valuation']['pe_ratio'] = market_cap / net_income
    if market_cap and total_equity > 0:
        result['valuation']['pb_ratio'] = market_cap / total_equity
    if enterprise_value and ebitda > 0:
        result['valuation']['ev_to_ebitda'] = enterprise_value / ebitda
    growth_rate = data.get('growth_rate')
    if market_cap and growth_rate:
        result['valuation']['peg_ratio'] = result['valuation'].get('pe_ratio') / growth_rate if result['valuation'].get('pe_ratio') else None

    return result


def build_report(data, ratios, recommendation):
    """Build the report content as structured data."""
    company = data.get('company_name', 'Company')
    fiscal_year = data.get('fiscal_year', 'N/A')
    currency = data.get('currency', 'USD')

    report = {
        'title': f'Financial Research Report: {company}',
        'date': datetime.now().strftime('%B %d, %Y'),
        'company': company,
        'fiscal_year': fiscal_year,
        'currency': currency,
        'executive_summary': {},
        'sections': [],
        'risks': [],
        'recommendation': recommendation or 'Hold',
    }

    # Executive Summary
    income = data.get('income_statement', {})
    report['executive_summary'] = {
        'revenue': abs(income.get('revenue', 0)),
        'ebitda': abs(income.get('ebitda', 0)),
        'net_income': abs(income.get('net_income', 0)),
        'total_assets': abs(data.get('balance_sheet', {}).get('total_assets', 0)),
    }

    # Section 1: Profitability
    profit = ratios.get('profitability', {})
    report['sections'].append({
        'title': 'Profitability Analysis',
        'ratios': profit,
        'benchmarks': {
            'gross_margin': '30-50% (typical)',
            'operating_margin': '10-20% (typical)',
            'net_margin': '5-15% (typical)',
            'roa': '5-15% (healthy)',
            'roe': '10-25% (healthy)',
        },
        'flags': []
    })
    if profit.get('net_margin', 0) < 0:
        report['sections'][-1]['flags'].append('ðŸ”´ Negative Net Margin â€” company is unprofitable')
    if profit.get('roa', 0) < 0.05:
        report['sections'][-1]['flags'].append('ðŸŸ¡ Low ROA (<5%) â€” assets not generating sufficient returns')

    # Section 2: Liquidity
    liq = ratios.get('liquidity', {})
    report['sections'].append({
        'title': 'Liquidity Analysis',
        'ratios': liq,
        'benchmarks': {
            'current_ratio': '>1.2x (healthy)',
            'quick_ratio': '>1.0x (healthy)',
        },
        'flags': []
    })
    if liq.get('current_ratio', 99) < 1.0:
        report['sections'][-1]['flags'].append('ðŸ”´ Current Ratio <1.0x â€” liquidity crisis risk')
    if liq.get('current_ratio', 0) < 1.2:
        report['sections'][-1]['flags'].append('ðŸŸ¡ Current Ratio <1.2x â€” tight liquidity')

    # Section 3: Leverage
    lev = ratios.get('leverage', {})
    report['sections'].append({
        'title': 'Leverage Analysis',
        'ratios': lev,
        'benchmarks': {
            'debt_to_equity': '<2.0x (healthy)',
            'debt_to_assets': '<0.5 (healthy)',
            'interest_coverage': '>2.5x (healthy)',
        },
        'flags': []
    })
    if lev.get('debt_to_equity', 0) > 2.0:
        report['sections'][-1]['flags'].append('ðŸ”´ Debt-to-Equity >2.0x â€” over-leveraged')
    if lev.get('interest_coverage', 99) < 2.5:
        report['sections'][-1]['flags'].append('ðŸ”´ Interest Coverage <2.5x â€” debt servicing risk')

    # Section 4: Efficiency
    eff = ratios.get('efficiency', {})
    report['sections'].append({
        'title': 'Efficiency Analysis',
        'ratios': eff,
        'benchmarks': {
            'debtor_days': '30-45 days (typical)',
            'creditor_days': '30-60 days (typical)',
            'inventory_days': '30-60 days (typical)',
            'cash_conversion_cycle': '<60 days (healthy)',
        },
        'flags': []
    })
    if eff.get('debtor_days', 0) > 60:
        report['sections'][-1]['flags'].append('ðŸ”´ Debtor Days >60 â€” slow collections')
    if eff.get('creditor_days', 0) > 90:
        report['sections'][-1]['flags'].append('ðŸ”´ Creditor Days >90 â€” supplier relationships at risk')
    if eff.get('cash_conversion_cycle', 0) > 90:
        report['sections'][-1]['flags'].append('ðŸ”´ Cash Conversion Cycle >90 days â€” capital tied up')

    # Section 5: Valuation
    val = ratios.get('valuation', {})
    if val:
        report['sections'].append({
            'title': 'Valuation Analysis',
            'ratios': val,
            'benchmarks': {
                'pe_ratio': '10-20x (typical)',
                'pb_ratio': '1-3x (typical)',
                'ev_to_ebitda': '8-12x (typical)',
            },
            'flags': []
        })
        if val.get('pe_ratio', 0) > 30:
            report['sections'][-1]['flags'].append('ðŸ”´ P/E >30x â€” potentially over-valued')
        if val.get('pe_ratio', 0) < 10 and val.get('pe_ratio', 0) > 0:
            report['sections'][-1]['flags'].append('ðŸŸ¡ P/E <10x â€” potentially under-valued')

    # Section 6: Risk Assessment
    all_flags = []
    for section in report['sections']:
        all_flags.extend(section.get('flags', []))
    report['risks'] = all_flags

    return report


def format_text_report(report):
    """Format as human-readable text."""
    lines = []
    r = report
    lines.append('=' * 60)
    lines.append(r['title'])
    lines.append(f"Date: {r['date']}")
    lines.append(f"Fiscal Year: {r['fiscal_year']} | Currency: {r['currency']}")
    lines.append('=' * 60)
    lines.append('')

    # Executive Summary
    lines.append('EXECUTIVE SUMMARY')
    lines.append('-' * 60)
    es = r['executive_summary']
    lines.append(f"  Revenue:      ${es['revenue']:,.0f}")
    lines.append(f"  EBITDA:       ${es['ebitda']:,.0f}")
    lines.append(f"  Net Income:   ${es['net_income']:,.0f}")
    lines.append(f"  Total Assets: ${es['total_assets']:,.0f}")
    lines.append('')
    lines.append(f"  Recommendation: {r['recommendation'].upper()}")
    lines.append('')

    # Sections
    for section in r['sections']:
        lines.append('')
        lines.append(section['title'].upper())
        lines.append('-' * 60)
        ratios = section.get('ratios', {})
        for k, v in ratios.items():
            if v is None:
                continue
            label = k.replace('_', ' ').title()
            if isinstance(v, float):
                if abs(v) < 10 and '.' in str(v):
                    lines.append(f"  {label:<25} {v:.2%}" if abs(v) < 1 else f"  {label:<25} {v:.2f}x")
                else:
                    lines.append(f"  {label:<25} {v:.2f}")
            else:
                lines.append(f"  {label:<25} {v}")
        if section.get('flags'):
            lines.append('')
            for flag in section['flags']:
                lines.append(f"  {flag}")
        lines.append('')

    # Disclaimer
    lines.append('=' * 60)
    lines.append('DISCLAIMER: This report is for informational purposes only.')
    lines.append('Not financial advice. Consult a qualified financial advisor before investing.')
    lines.append('=' * 60)

    return '\n'.join(lines)


def write_word_report(report, output_path):
    """Generate a professional Word document."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN
    except ImportError:
        print("Error: python-docx is required. Install with: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document()
    r = report

    # Title
    title = doc.add_heading(r['title'])
    title.alignment = WD_ALIGN.CENTER

    doc.add_paragraph(f"Date: {r['date']}")
    doc.add_paragraph(f"Fiscal Year: {r['fiscal_year']} | Currency: {r['currency']}")
    doc.add_paragraph('')

    # Executive Summary
    doc.add_heading('Executive Summary', level=2)
    es = r['executive_summary']
    for label, key in [('Revenue', 'revenue'), ('EBITDA', 'ebitda'), ('Net Income', 'net_income'), ('Total Assets', 'total_assets')]:
        doc.add_paragraph(f"{label}: ${es[key]:,.0f}" if isinstance(es[key], (int, float)) else f"{label}: {es[key]}")

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run(f"Recommendation: {r['recommendation'].upper()}").bold = True

    # Sections
    for section in r['sections']:
        doc.add_heading(section['title'], level=2)
        ratios = section.get('ratios', {})
        for k, v in ratios.items():
            if v is None:
                continue
            label = k.replace('_', ' ').title()
            if isinstance(v, float):
                if 0 < abs(v) < 1:
                    doc.add_paragraph(f"{label}: {v:.2%}")
                else:
                    doc.add_paragraph(f"{label}: {v:.2f}")
            else:
                doc.add_paragraph(f"{label}: {v}")
        if section.get('flags'):
            doc.add_paragraph('Flags:', style='Intense Quote')
            for flag in section['flags']:
                doc.add_paragraph(flag, style='List Bullet')

    # Disclaimer
    doc.add_paragraph('')
    doc.add_paragraph('DISCLAIMER: This report is for informational purposes only. Not financial advice.', style='Intense Quote')
    doc.add_paragraph('Consult a qualified financial advisor before making investment decisions.', style='Intense Quote')

    doc.save(output_path)
    print(f"Saved Word report: {output_path}")


def main():
    parser = argparse.ArgumentParser(description='Generate a professional financial research report.')
    parser.add_argument('input_json', help='Path to financial data JSON file')
    parser.add_argument('--output', dest='output', default=None,
                        help='Output .docx file path')
    parser.add_argument('--recommendation', dest='rec', default=None,
                        choices=['Buy', 'Hold', 'Sell'],
                        help='Investment recommendation (Buy/Hold/Sell)')
    parser.add_argument('--format', dest='fmt', choices=['json', 'text', 'docx'],
                        default='docx', help='Output format (default: docx)')
    args = parser.parse_args()

    try:
        with open(args.input_json, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"Error: File not found: {args.input_json}", file=sys.stderr)
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON: {e}", file=sys.stderr)
        sys.exit(1)

    ratios = calculate_all_ratios(data)
    report = build_report(data, ratios, args.rec)

    if args.fmt == 'json':
        output = json.dumps(report, indent=2, default=str)
        if args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                f.write(output)
            print(f"Saved JSON: {args.output}")
        else:
            print(output)
    elif args.fmt == 'text':
        output = format_text_report(report)
        if args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                f.write(output)
            print(f"Saved text report: {args.output}")
        else:
            print(output)
    else:  # docx
        output_path = args.output
        if not output_path:
            safe_name = data.get('company_name', 'Company').replace(' ', '_')
            output_path = f"{safe_name}_financial_report.docx"
        write_word_report(report, output_path)


if __name__ == '__main__':
    main()
